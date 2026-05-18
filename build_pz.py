from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("C:/diploma/ПЗ_NEO-SYNC_Асылбеков.docx")
TOPIC = (
    "Проектирование и разработка веб-системы записи на обслуживание "
    "для киберпанкового сервисного центра компьютерной техники «NEO-SYNC»"
)


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_border(cell, color="777777", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def setup(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(1)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.25)
    sec.different_first_page_header_footer = True
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for name, size in (("Heading 1", 16), ("Heading 2", 15), ("Heading 3", 14)):
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = True
        st.paragraph_format.first_line_indent = Cm(0)
        st.paragraph_format.line_spacing = 1.5
        st.paragraph_format.space_before = Pt(10)
        st.paragraph_format.space_after = Pt(6)
    doc.styles["Title"].font.size = Pt(16)


def p(doc, text="", bold=False, center=False, no_indent=False):
    par = doc.add_paragraph(style="Normal")
    if center:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if no_indent:
        par.paragraph_format.first_line_indent = Cm(0)
    run = par.add_run(text)
    run.bold = bold
    return par


def h(doc, level, text):
    par = doc.add_paragraph(style=f"Heading {level}")
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    par.add_run(text.upper() if level == 1 else text)
    return par


def caption(doc, text):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.first_line_indent = Cm(0)
    par.paragraph_format.line_spacing = 1.0
    run = par.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    return par


def table(doc, headers, rows, widths):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    for i, head in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.width = Cm(widths[i])
        set_shading(cell, "EDEDED")
        set_border(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        par = cell.paragraphs[0]
        par.paragraph_format.first_line_indent = Cm(0)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = par.add_run(head)
        run.bold = True
        run.font.size = Pt(12)
    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].width = Cm(widths[i])
            set_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            par = cells[i].paragraphs[0]
            par.paragraph_format.first_line_indent = Cm(0)
            run = par.add_run(str(val))
            run.font.size = Pt(12)
    return tbl


def code(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_shading(cell, "F2F2F2")
    set_border(cell, color="999999")
    par = cell.paragraphs[0]
    par.paragraph_format.first_line_indent = Cm(0)
    run = par.add_run(text)
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(10)


def screenshot(doc, number, title, note):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    tbl.rows[0].height = Cm(10.5)
    cell.width = Cm(10.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_shading(cell, "F7F7F7")
    set_border(cell, color="333333", size="18")
    par = cell.paragraphs[0]
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.paragraph_format.first_line_indent = Cm(0)
    r = par.add_run("МЕСТО ДЛЯ СКРИНШОТА")
    r.bold = True
    r.font.size = Pt(14)
    par2 = cell.add_paragraph()
    par2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par2.paragraph_format.first_line_indent = Cm(0)
    run = par2.add_run(note)
    run.font.size = Pt(11)
    caption(doc, f"Рисунок {number} - {title}")


def cover(doc):
    lines = [
        "КЫРГЫЗ РЕСПУБЛИКАСЫНЫН БИЛИМ БЕРҮҮ ЖАНА ИЛИМ МИНИСТРЛИГИ",
        "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ КЫРГЫЗСКОЙ РЕСПУБЛИКИ",
        "",
        "И. Раззаков атындагы КЫРГЫЗ МАМЛЕКЕТТИК ТЕХНИКАЛЫК УНИВЕРСИТЕТИ",
        "КЫРГЫЗСКИЙ ГОСУДАРСТВЕННЫЙ ТЕХНИЧЕСКИЙ УНИВЕРСИТЕТ имени И. Раззакова",
        "",
        "КЕСИПТИК ОРТО БИЛИМ БЕРҮҮ (КОЛЛЕДЖ)",
        "СРЕДНЕЕ ПРОФЕССИОНАЛЬНОЕ ОБРАЗОВАНИЕ (КОЛЛЕДЖ)",
    ]
    for line in lines:
        p(doc, line, bold=True, center=True, no_indent=True)
    for _ in range(4):
        p(doc, "", no_indent=True)
    p(doc, "ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", bold=True, center=True, no_indent=True)
    p(doc, "к выпускной квалификационной работе", center=True, no_indent=True)
    p(doc, f"на тему: «{TOPIC}»", bold=True, center=True, no_indent=True)
    for _ in range(3):
        p(doc, "", no_indent=True)
    rows = [
        ("Студент группы", "__________  Асылбеков __________________  _________"),
        ("Образовательная программа", "230111 - Программирование в компьютерных системах"),
        ("Форма обучения", "очная"),
        ("Руководитель", "______________________________  _________"),
        ("Консультант", "______________________________  _________"),
        ("Председатель ПЦК", "______________________________  _________"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for left, right in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate((left, right)):
            cells[i].width = Cm(5 if i == 0 else 11)
            par = cells[i].paragraphs[0]
            par.paragraph_format.first_line_indent = Cm(0)
            run = par.add_run(value)
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)
    p(doc, "", no_indent=True)
    p(doc, "Бишкек - 2026", center=True, no_indent=True)
    doc.add_page_break()


def assignment(doc):
    h(doc, 1, "Задание на выпускную квалификационную работу")
    p(doc, "Студент группы __________ Асылбеков ________________________________.")
    p(doc, f"Тема ВКР: «{TOPIC}».")
    p(doc, "Исходные данные к работе: MVP веб-проекта NEO-SYNC, реализованный на FastAPI, SQLite и чистом HTML/CSS/JavaScript; требования к расширению функциональности сайта, проектированию базы данных до третьей нормальной формы и подготовке архитектуры серверной части по layered-подходу.")
    table(doc, ["№", "Содержание расчетно-пояснительной записки", "Объем, %", "Срок"], [
        ("1", "Анализ предметной области и постановка задачи", "15", "____"),
        ("2", "Проектирование архитектуры веб-приложения", "15", "____"),
        ("3", "Проектирование нормализованной базы данных", "20", "____"),
        ("4", "Реализация backend-части и API", "20", "____"),
        ("5", "Реализация пользовательского интерфейса", "15", "____"),
        ("6", "Тестирование, безопасность и оценка результата", "10", "____"),
        ("7", "Список источников и приложения", "5", "____"),
    ], [1, 10, 2.5, 2.5])
    p(doc, "Графическая часть: схема базы данных, диаграмма архитектуры приложения, экранные формы сайта, скриншоты DBeaver, скриншоты документации API и админ-панели.")
    p(doc, "Дата выдачи задания: «___» __________ 2026 г.")
    p(doc, "Руководитель ВКР: ________________________________")
    p(doc, "Задание получил: ________________________________")
    doc.add_page_break()


def contents(doc):
    h(doc, 1, "Содержание")
    entries = [
        ("Введение", "5"),
        ("Глава 1. Анализ и проектирование веб-системы NEO-SYNC", "7"),
        ("1.1 Анализ предметной области", "7"),
        ("1.2 Обзор используемых технологий", "10"),
        ("1.3 Требования к системе", "13"),
        ("1.4 Нормализация данных и требования 1НФ, 2НФ, 3НФ", "16"),
        ("Глава 2. Реализация веб-системы и базы данных", "20"),
        ("2.1 Layered-архитектура backend-приложения", "20"),
        ("2.2 Проектирование структуры базы данных", "23"),
        ("2.3 Создание базы данных в SQLite и DBeaver", "27"),
        ("2.4 Реализация API и валидации данных", "32"),
        ("2.5 Реализация пользовательского интерфейса", "36"),
        ("2.6 Тестирование и оценка результата", "41"),
        ("Заключение", "44"),
        ("Список использованных источников", "46"),
        ("Приложения", "48"),
    ]
    for title, page in entries:
        par = p(doc, "", no_indent=True)
        par.paragraph_format.tab_stops.add_tab_stop(Cm(16))
        par.add_run(title)
        par.add_run("\t" + page)
    doc.add_page_break()


def add_many(doc, paragraphs):
    for text in paragraphs:
        p(doc, text)


def intro(doc):
    h(doc, 1, "Введение")
    add_many(doc, [
        "Современные предприятия сферы обслуживания компьютерной и офисной техники нуждаются в цифровых инструментах, которые позволяют быстро принимать заявки, хранить информацию об услугах, контролировать расписание работ и обеспечивать удобную коммуникацию с клиентами.",
        "В рамках данной выпускной квалификационной работы рассматривается проектирование и разработка веб-системы «NEO-SYNC». Проект представляет собой киберпанковый сервисный центр компьютерной техники, где пользователь может просмотреть каталог услуг, выбрать конкретную услугу, отправить заявку на обслуживание, а администратор может просматривать и обрабатывать поступившие записи.",
        "Актуальность темы обусловлена тем, что веб-приложения для онлайн-записи являются распространенным инструментом автоматизации обслуживания клиентов. Для таких систем особенно важны корректная структура базы данных, целостность связей между сущностями, валидация входных данных, понятная архитектура backend-части и возможность дальнейшего расширения функциональности.",
        "Целью работы является разработка веб-системы записи на обслуживание для сервисного центра компьютерной техники «NEO-SYNC» с нормализованной базой данных SQLite и backend-частью на FastAPI, построенной по layered-архитектуре.",
        "Для достижения поставленной цели необходимо решить следующие задачи: провести анализ предметной области; определить основные сущности системы; спроектировать базу данных, соответствующую первой, второй и третьей нормальным формам; реализовать структуру backend-приложения по слоям Domain, Application, Infrastructure и Api; настроить работу с SQLite через репозитории; реализовать Pydantic-схемы для валидации данных; подготовить API для заявок, каталога услуг, обратной связи и администрирования.",
        "Объектом исследования является процесс автоматизации записи клиентов на услуги сервисного центра компьютерной техники. Предметом исследования является разработка структуры базы данных и серверной части веб-приложения для приема, хранения и обработки заявок.",
        "Практическая значимость работы заключается в создании MVP-системы, которую можно расширять: добавлять авторизацию администратора, личный кабинет клиента, управление статусами заявок, портфолио выполненных проектов, аналитику популярности услуг и интеграцию с внешними сервисами уведомлений.",
    ])
    screenshot(doc, 1, "Главная страница веб-системы NEO-SYNC", "Скриншот браузера: первый экран сайта с логотипом NEO-SYNC, навигацией и киберпанковым оформлением.")
    doc.add_page_break()


def chapter1(doc):
    h(doc, 1, "Глава 1. Анализ и проектирование веб-системы NEO-SYNC")
    h(doc, 2, "1.1 Анализ предметной области")
    add_many(doc, [
        "Предметная область проекта связана с деятельностью сервисного центра компьютерной техники, которая предоставляет клиентам различные услуги: диагностику, настройку производительности, установку защитных систем, ремонт и консультации. Клиент должен иметь возможность ознакомиться с перечнем услуг, выбрать подходящую услугу и отправить заявку на конкретную дату и время.",
        "В ручном варианте учет заявок может выполняться в таблицах, мессенджерах или блокнотах. Такой подход плохо масштабируется: отсутствует единая структура данных, сложно контролировать статус заявки, не обеспечивается целостность информации, а повторяющиеся сведения об услугах и категориях приводят к избыточности.",
        "Основными пользователями системы являются клиент и администратор. Клиент просматривает сайт, выбирает услугу, заполняет форму записи и может отправить сообщение через контактную форму. Администратор просматривает заявки, контролирует их наличие и в дальнейшем сможет менять статусы, подтверждать запись и формировать отчеты.",
    ])
    table(doc, ["Роль", "Функции в системе"], [
        ("Клиент", "Просмотр каталога услуг, выбор услуги, отправка заявки, отправка сообщения через форму контактов."),
        ("Администратор", "Просмотр заявок, удаление ошибочных записей, дальнейшее изменение статусов и контроль расписания."),
        ("Система", "Валидация данных, сохранение информации в SQLite, обеспечение целостности связей, отдача статических файлов."),
    ], [4, 12])
    caption(doc, "Таблица 1 - Роли пользователей веб-системы")
    h(doc, 2, "1.2 Обзор используемых технологий")
    add_many(doc, [
        "Для реализации backend-части выбран FastAPI. Данный фреймворк позволяет быстро создавать REST API, поддерживает автоматическую генерацию OpenAPI-документации, интегрируется с Pydantic для строгой валидации данных и подходит для учебных и производственных веб-сервисов.",
        "В качестве базы данных используется SQLite. Для дипломного MVP SQLite является рациональным выбором, так как не требует отдельного сервера, хранит данные в одном файле, поддерживает SQL, внешние ключи, индексы, ограничения целостности и транзакции. Для просмотра структуры базы данных и работы с таблицами используется DBeaver.",
        "Frontend реализован на чистом HTML, CSS и JavaScript. Такой подход позволяет показать базовую механику взаимодействия клиента с API без дополнительной сложности фреймворков. Статические файлы отдаются самим FastAPI-приложением через StaticFiles, что устраняет жесткую привязку frontend-части к абсолютному адресу сервера.",
    ])
    table(doc, ["Компонент", "Технология", "Назначение"], [
        ("Backend", "FastAPI", "Создание REST API, маршрутизация, обработка HTTP-запросов."),
        ("Валидация", "Pydantic", "Проверка входных данных и формирование DTO-схем."),
        ("База данных", "SQLite", "Хранение категорий, услуг, заявок, портфолио и обратной связи."),
        ("Редактор БД", "DBeaver", "Просмотр таблиц, ER-диаграммы, выполнение SQL-запросов."),
        ("Frontend", "HTML/CSS/JavaScript", "Пользовательский интерфейс сайта и админ-панели."),
    ], [3.5, 4, 8.5])
    caption(doc, "Таблица 2 - Используемые технологии")
    screenshot(doc, 2, "Окно DBeaver с подключенной базой service.db", "Скриншот DBeaver: слева подключение SQLite service.db, раскрыт список таблиц базы данных.")
    h(doc, 2, "1.3 Требования к системе")
    p(doc, "Функциональные требования определяют, какие действия должна выполнять система. Для NEO-SYNC выделены требования к публичной части сайта, административной части, API и базе данных.")
    table(doc, ["№", "Требование", "Описание"], [
        ("F1", "Каталог услуг", "Система должна хранить категории и услуги, отображать их на сайте и отдавать через API."),
        ("F2", "Запись на обслуживание", "Клиент должен иметь возможность выбрать услугу, дату и отправить заявку."),
        ("F3", "Админ-панель", "Администратор должен просматривать заявки и удалять ошибочные записи."),
        ("F4", "Форма контактов", "Пользователь должен отправлять сообщение с контактными данными."),
        ("F5", "Портфолио", "Система должна поддерживать хранение проектов и связь проектов с услугами."),
        ("F6", "Статические файлы", "Backend должен отдавать HTML, CSS и JavaScript без отдельного frontend-сервера."),
    ], [1.5, 4, 10.5])
    caption(doc, "Таблица 3 - Функциональные требования")
    p(doc, "Нефункциональные требования включают простоту развертывания, модульность, читаемость кода, возможность дальнейшего расширения, базовую безопасность административных эндпоинтов и сохранение целостности данных средствами СУБД.")
    h(doc, 2, "1.4 Нормализация данных и требования 1НФ, 2НФ, 3НФ")
    add_many(doc, [
        "Первоначальная версия проекта содержала одну таблицу maintenance_requests, где вместе хранились имя клиента, устройство, дата обслуживания и текстовое название услуги. Такая структура подходит для MVP, но при расширении системы приводит к избыточности.",
        "Для устранения избыточности база данных была перепроектирована. Услуги вынесены в отдельную таблицу Services, категории - в Categories, заявки - в Bookings, проекты портфолио - в Portfolio_Projects, а связь «многие-ко-многим» между проектами и услугами реализована через Project_Services.",
        "Первая нормальная форма обеспечивается тем, что все поля таблиц атомарны: в одной ячейке хранится одно значение, повторяющиеся группы отсутствуют. Вторая нормальная форма выполняется, так как неключевые атрибуты зависят от полного первичного ключа. Третья нормальная форма достигается за счет отсутствия транзитивных зависимостей.",
    ])
    table(doc, ["Нормальная форма", "Как обеспечена в проекте"], [
        ("1НФ", "Все атрибуты атомарны; списки услуг проекта вынесены в отдельную связующую таблицу."),
        ("2НФ", "Неключевые поля зависят от полного первичного ключа; составной ключ используется только в Project_Services."),
        ("3НФ", "Справочные данные об услугах и категориях не дублируются в заявках и проектах."),
    ], [4, 12])
    caption(doc, "Таблица 4 - Соответствие базы данных нормальным формам")
    screenshot(doc, 3, "ER-диаграмма базы данных в DBeaver", "Скриншот DBeaver: ER Diagram с таблицами Categories, Services, Bookings, Feedbacks, Portfolio_Projects, Project_Services и связями FK.")
    doc.add_page_break()


def chapter2(doc):
    h(doc, 1, "Глава 2. Реализация веб-системы и базы данных")
    h(doc, 2, "2.1 Layered-архитектура backend-приложения")
    p(doc, "Backend-часть проекта была переработана в соответствии с layered-архитектурой. Такой подход разделяет ответственность между слоями и делает код более понятным для сопровождения.")
    table(doc, ["Слой", "Папка", "Назначение"], [
        ("Domain", "app/domain", "Доменные сущности, перечисления, бизнес-исключения."),
        ("Application", "app/application", "Pydantic-схемы, интерфейсы репозиториев, сервисы бизнес-сценариев."),
        ("Infrastructure", "app/infrastructure", "Подключение к SQLite, создание таблиц, репозитории."),
        ("Api", "app/api", "FastAPI-приложение, маршруты, зависимости, CORS, StaticFiles."),
    ], [3, 4, 9])
    caption(doc, "Таблица 5 - Слои backend-архитектуры")
    code(doc, "from app.api.app import create_app\n\napp = create_app()")
    screenshot(doc, 4, "Структура проекта в редакторе кода", "Скриншот IDE или проводника: папки app/domain, app/application, app/infrastructure, app/api, app/static.")
    h(doc, 2, "2.2 Проектирование структуры базы данных")
    p(doc, "Новая структура базы данных включает шесть основных таблиц. Таблицы Categories и Services формируют каталог услуг. Таблица Bookings хранит заявки клиентов и ссылается на конкретную услугу. Таблица Feedbacks хранит сообщения из формы контактов.")
    table(doc, ["Таблица", "Основные поля", "Назначение"], [
        ("Categories", "id, name, slug", "Справочник категорий услуг."),
        ("Services", "id, category_id, title, description, price", "Каталог услуг с ценой и описанием."),
        ("Portfolio_Projects", "id, title, description, client_name, completion_date", "Кейсы и выполненные проекты."),
        ("Project_Services", "project_id, service_id", "Связь многие-ко-многим между проектами и услугами."),
        ("Bookings", "id, client_name, device, service_id, service_date, status, created_at", "Заявки клиентов на обслуживание."),
        ("Feedbacks", "id, name, email_or_phone, message, created_at", "Сообщения из формы контактов."),
    ], [4, 6, 6])
    caption(doc, "Таблица 6 - Основные таблицы базы данных")
    code(doc, "CREATE TABLE IF NOT EXISTS Bookings (\n    id INTEGER PRIMARY KEY AUTOINCREMENT,\n    client_name TEXT NOT NULL,\n    device TEXT NOT NULL,\n    service_id INTEGER NOT NULL,\n    service_date TEXT NOT NULL,\n    status TEXT NOT NULL DEFAULT 'new',\n    created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,\n    FOREIGN KEY (service_id) REFERENCES Services(id)\n);")
    screenshot(doc, 5, "Список таблиц базы данных в DBeaver", "Скриншот DBeaver: раскрытый узел Tables с таблицами Categories, Services, Portfolio_Projects, Project_Services, Bookings, Feedbacks.")
    h(doc, 2, "2.3 Создание базы данных в SQLite и DBeaver")
    add_many(doc, [
        "Файл базы данных проекта называется service.db и расположен в корневой директории проекта. При запуске backend-приложения вызывается функция init_db, которая создает таблицы, если они отсутствуют.",
        "Для визуального контроля структуры базы используется DBeaver. В DBeaver создается подключение SQLite к файлу service.db. После подключения можно открыть список таблиц, просмотреть их поля, типы данных, внешние ключи, индексы и содержимое записей.",
        "Особое внимание уделяется проверке внешних ключей. SQLite требует включения PRAGMA foreign_keys = ON для корректного контроля ссылочной целостности. В коде подключения этот режим включается при создании соединения.",
    ])
    code(doc, 'conn = sqlite3.connect(path)\nconn.row_factory = sqlite3.Row\nconn.execute("PRAGMA foreign_keys = ON")\nconn.execute("PRAGMA journal_mode = WAL")')
    screenshot(doc, 6, "Подключение SQLite в DBeaver", "Скриншот DBeaver: окно создания или настройки SQLite-подключения к файлу C:/diploma/service.db.")
    screenshot(doc, 7, "Просмотр данных таблицы Services", "Скриншот DBeaver: вкладка Data таблицы Services с услугами диагностики ноутбука/ПК, чистки и настройки принтера.")
    screenshot(doc, 8, "Просмотр внешних ключей таблицы Bookings", "Скриншот DBeaver: свойства таблицы Bookings, раздел Foreign Keys, связь service_id -> Services.id.")
    h(doc, 2, "2.4 Реализация API и валидации данных")
    p(doc, "API реализован через FastAPI-роутеры. Публичные маршруты имеют префикс /api, административные маршруты - /admin. Для будущей авторизации подготовлена dependency require_admin, которая проверяет заголовок X-Admin-Token.")
    table(doc, ["Метод", "URL", "Назначение"], [
        ("GET", "/api/services", "Получить каталог услуг."),
        ("POST", "/api/bookings", "Создать заявку на обслуживание."),
        ("POST", "/api/feedbacks", "Отправить сообщение через форму контактов."),
        ("GET", "/admin/bookings", "Получить список заявок администратора."),
        ("DELETE", "/admin/bookings/{booking_id}", "Удалить заявку; при отсутствии ID возвращается 404."),
    ], [2, 5, 9])
    caption(doc, "Таблица 7 - Основные API-эндпоинты")
    code(doc, 'class BookingCreate(BaseModel):\n    client_name: ShortText\n    device: ShortText\n    service_id: int = Field(gt=0)\n    service_date: datetime')
    screenshot(doc, 9, "Swagger UI FastAPI", "Скриншот браузера: страница /docs с маршрутами /api/services, /api/bookings, /api/feedbacks и /admin/bookings.")
    screenshot(doc, 10, "Пример ошибки 404 при удалении несуществующей заявки", "Скриншот Swagger UI или браузера: DELETE /admin/bookings/{id}, ответ 404 Booking with id=... was not found.")
    h(doc, 2, "2.5 Реализация пользовательского интерфейса")
    p(doc, "Пользовательский интерфейс реализован в виде статических HTML-страниц. Главная страница содержит навигацию, каталог услуг, форму записи и форму обратной связи. Административная страница отображает список заявок и позволяет удалить ошибочные записи.")
    screenshot(doc, 11, "Каталог услуг на сайте", "Скриншот браузера: блок «Каталог услуг» с карточками услуг, загруженными через GET /api/services.")
    screenshot(doc, 12, "Форма записи на обслуживание", "Скриншот браузера: форма с полями имя, устройство, услуга, дата и кнопкой отправки заявки.")
    screenshot(doc, 13, "Форма контактов", "Скриншот браузера: форма контактов с полями имя, email/телефон и сообщение.")
    screenshot(doc, 14, "Административная панель", "Скриншот браузера: страница admin.html со списком заявок и кнопкой удаления.")
    h(doc, 2, "2.6 Тестирование и оценка результата")
    p(doc, "Проверка проекта выполнялась на нескольких уровнях. На уровне структуры кода проверяется импорт приложения и компиляция Python-модулей. На уровне базы данных проверяется создание таблиц, наличие seed-данных и отсутствие ошибок внешних ключей.")
    table(doc, ["Проверка", "Ожидаемый результат", "Фактический результат"], [
        ("Импорт main.py", "Приложение FastAPI создается без ошибки.", "Успешно."),
        ("Инициализация БД", "Создаются таблицы Categories, Services, Bookings и др.", "Успешно."),
        ("PRAGMA foreign_key_check", "Список ошибок пустой.", "Ошибок нет."),
        ("GET /api/services", "Возвращается список услуг.", "Маршрут зарегистрирован."),
        ("DELETE отсутствующей заявки", "Возвращается ошибка 404.", "Реализовано через EntityNotFoundError."),
        ("GET /admin/bookings без токена", "Возвращается 401.", "Реализовано через require_admin."),
    ], [5, 6, 5])
    caption(doc, "Таблица 8 - Результаты проверки")
    screenshot(doc, 15, "Результат SQL-запроса проверки внешних ключей", "Скриншот DBeaver: выполнение PRAGMA foreign_key_check; пустой результат или отсутствие строк ошибок.")
    screenshot(doc, 16, "Созданная заявка в таблице Bookings", "Скриншот DBeaver: таблица Bookings после отправки формы записи, видна новая запись со статусом new.")
    doc.add_page_break()


def conclusion(doc):
    h(doc, 1, "Заключение")
    add_many(doc, [
        "В ходе выполнения выпускной квалификационной работы была спроектирована и реализована веб-система «NEO-SYNC» для записи клиентов на обслуживание сервисного центра компьютерной техники. Проект был переработан из первоначального MVP в более структурированную систему с нормализованной базой данных и layered-архитектурой backend-части.",
        "Была выполнена декомпозиция предметной области на сущности: категории услуг, услуги, проекты портфолио, связи проектов с услугами, заявки на обслуживание и сообщения обратной связи. На основе этих сущностей разработана структура базы данных SQLite, соответствующая требованиям первой, второй и третьей нормальных форм.",
        "Backend-приложение реализовано на FastAPI и разделено на слои Domain, Application, Infrastructure и Api. Такой подход улучшает читаемость проекта, снижает связанность компонентов и подготавливает систему к дальнейшему расширению.",
        "Практическая часть показала, что выбранный стек FastAPI, SQLite, DBeaver и чистый HTML/CSS/JavaScript подходит для разработки учебного дипломного веб-проекта. Система может быть расширена за счет личного кабинета клиента, управления статусами заявок, загрузки реальных изображений портфолио, отчетов по популярности услуг, интеграции с email/SMS-уведомлениями и более строгой системы прав доступа.",
    ])
    doc.add_page_break()


def bibliography(doc):
    h(doc, 1, "Список использованных источников")
    sources = [
        "FastAPI Documentation. URL: https://fastapi.tiangolo.com/",
        "Pydantic Documentation. URL: https://docs.pydantic.dev/",
        "SQLite Documentation. URL: https://www.sqlite.org/docs.html",
        "DBeaver Documentation. URL: https://dbeaver.com/docs/dbeaver/",
        "MDN Web Docs: HTML. URL: https://developer.mozilla.org/ru/docs/Web/HTML",
        "MDN Web Docs: CSS. URL: https://developer.mozilla.org/ru/docs/Web/CSS",
        "MDN Web Docs: Fetch API. URL: https://developer.mozilla.org/ru/docs/Web/API/Fetch_API",
        "Starlette StaticFiles Documentation. URL: https://www.starlette.io/staticfiles/",
        "ГОСТ 7.32-2017. Отчет о научно-исследовательской работе. Структура и правила оформления.",
        "ГОСТ 2.105-2019. Единая система конструкторской документации. Общие требования к текстовым документам.",
        "Кузнецов С.Д. Базы данных: модели и языки. - М.: Интернет-Университет Информационных Технологий.",
        "Карпова Т.С. Базы данных: модели, разработка, реализация. - СПб.: Питер.",
    ]
    for i, source in enumerate(sources, 1):
        p(doc, f"{i}. {source}", no_indent=True)
    doc.add_page_break()


def appendices(doc):
    h(doc, 1, "Приложения")
    h(doc, 2, "Приложение А. SQL-скрипт создания таблиц")
    code(doc, "CREATE TABLE Categories (...);\nCREATE TABLE Services (...);\nCREATE TABLE Portfolio_Projects (...);\nCREATE TABLE Project_Services (...);\nCREATE TABLE Bookings (...);\nCREATE TABLE Feedbacks (...);\nCREATE INDEX idx_services_category_id ON Services(category_id);\nCREATE INDEX idx_bookings_service_date ON Bookings(service_date);")
    h(doc, 2, "Приложение Б. Перечень скриншотов для вставки")
    table(doc, ["№ рисунка", "Что вставить вместо квадратного плейсхолдера"], [
        ("1", "Главная страница NEO-SYNC в браузере."),
        ("2", "DBeaver: подключенная база service.db."),
        ("3", "DBeaver: ER-диаграмма базы данных."),
        ("4", "Структура проекта в редакторе кода."),
        ("5", "DBeaver: список таблиц."),
        ("6", "DBeaver: окно подключения SQLite."),
        ("7", "DBeaver: данные таблицы Services."),
        ("8", "DBeaver: внешний ключ Bookings.service_id."),
        ("9", "Swagger UI FastAPI."),
        ("10", "Ошибка 404 при удалении отсутствующей заявки."),
        ("11", "Каталог услуг на сайте."),
        ("12", "Форма записи на обслуживание."),
        ("13", "Форма контактов."),
        ("14", "Административная панель."),
        ("15", "DBeaver: PRAGMA foreign_key_check."),
        ("16", "DBeaver: новая запись в Bookings."),
    ], [3, 13])
    caption(doc, "Таблица 9 - Список необходимых скриншотов")


def main():
    doc = Document()
    setup(doc)
    cover(doc)
    assignment(doc)
    contents(doc)
    intro(doc)
    chapter1(doc)
    chapter2(doc)
    conclusion(doc)
    bibliography(doc)
    appendices(doc)
    for section in doc.sections:
        add_page_number(section.footer.paragraphs[0])
    doc.core_properties.title = "Пояснительная записка NEO-SYNC"
    doc.core_properties.subject = TOPIC
    doc.core_properties.author = "Асылбеков"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
